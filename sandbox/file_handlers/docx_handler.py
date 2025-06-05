"""
DOCX file handler with Streamlit optimization.
Supports comprehensive Word document content extraction including text, tables, and structure.
"""

import streamlit as st
import io
import time
from typing import Dict, Any, List, Optional
from .base_handler import BaseFileHandler, ExtractedContent, FileType

# Try to import python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DOCXHandler(BaseFileHandler):
    """Streamlit-optimized DOCX content handler"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['docx']
        self.handler_name = "DOCXHandler"
        self.docx_available = DOCX_AVAILABLE
        
        if not DOCX_AVAILABLE:
            st.warning("⚠️ python-docx not available. DOCX processing will be limited.")
    
    def validate_file(self, uploaded_file) -> bool:
        """Validate DOCX file"""
        try:
            # Check file extension
            if not uploaded_file.name.lower().endswith('.docx'):
                return False
            
            # Check file size (max 10MB for Streamlit)
            if not self._validate_file_size(uploaded_file, max_size_mb=10):
                return False
            
            # Try to read first few bytes to validate DOCX format (ZIP signature)
            current_pos = uploaded_file.tell()
            file_bytes = uploaded_file.read(4)
            uploaded_file.seek(current_pos)  # Reset file pointer
            
            return file_bytes.startswith(b'PK\x03\x04')
        
        except Exception:
            return False
    
    def extract_content(self, uploaded_file) -> ExtractedContent:
        """Extract comprehensive content from DOCX"""
        start_time = time.time()
        
        try:
            if self.docx_available:
                # Read file bytes
                file_bytes = uploaded_file.read()
                uploaded_file.seek(0)  # Reset for potential reuse
                
                # Use cached extraction for performance
                content = self._extract_docx_cached(file_bytes, uploaded_file.name)
            else:
                # Fallback when python-docx not available
                content = self._extract_docx_fallback(uploaded_file)
            
            # Add processing time
            content.processing_time = time.time() - start_time
            
            return content
        
        except Exception as e:
            return self._create_error_content(
                f"DOCX processing error: {str(e)}", 
                uploaded_file.name
            )
    
    @st.cache_data
    def _extract_docx_cached(_self, file_bytes: bytes, filename: str) -> ExtractedContent:
        """Cached DOCX content extraction"""
        
        if not DOCX_AVAILABLE:
            return ExtractedContent(
                text="",
                metadata={'filename': filename, 'error': True},
                tables=[],
                images=[],
                structure={},
                errors=["python-docx library not available"]
            )
        
        try:
            # Create file-like object
            docx_file = io.BytesIO(file_bytes)
            
            # Open document
            doc = Document(docx_file)
            
            # Extract text content
            text_content = _self._extract_text_content(doc)
            
            # Extract tables
            tables = _self._extract_tables(doc)
            
            # Extract document structure
            structure = _self._analyze_document_structure(doc)
            
            # Extract metadata
            metadata = _self._extract_metadata(doc, filename, len(file_bytes))
            
            # Detect images/shapes (basic detection)
            images = _self._detect_images(doc)
            
            return ExtractedContent(
                text=text_content,
                metadata=metadata,
                tables=tables,
                images=images,
                structure=structure
            )
        
        except Exception as e:
            return ExtractedContent(
                text="",
                metadata={'filename': filename, 'error': True},
                tables=[],
                images=[],
                structure={},
                errors=[f"DOCX extraction failed: {str(e)}"]
            )
    
    def _extract_docx_fallback(self, uploaded_file) -> ExtractedContent:
        """Fallback extraction when python-docx not available"""
        
        return ExtractedContent(
            text=f"DOCX file detected: {uploaded_file.name}",
            metadata={
                'filename': uploaded_file.name,
                'size_bytes': uploaded_file.size,
                'extraction_method': 'fallback',
                'note': 'python-docx library required for full DOCX processing'
            },
            tables=[],
            images=[],
            structure={'extraction_method': 'fallback'},
            warnings=["python-docx library not available - limited processing"]
        )
    
    def _extract_text_content(self, doc) -> str:
        """Extract all text content from document"""
        
        text_parts = []
        
        try:
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables (separate from table structure)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_parts.append('\n'.join(table_text))
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Header text
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[HEADER] {paragraph.text}")
                
                # Footer text
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[FOOTER] {paragraph.text}")
        
        except Exception as e:
            text_parts.append(f"Error extracting text: {str(e)}")
        
        return '\n'.join(text_parts)
    
    def _extract_tables(self, doc) -> List[Dict[str, Any]]:
        """Extract table data with structure"""
        
        tables = []
        
        try:
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_index': table_idx + 1,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0,
                        'has_header': self._detect_table_header(table_data),
                        'table_type': self._classify_table_type(table_data)
                    })
        
        except Exception as e:
            tables.append({
                'error': f"Table extraction failed: {str(e)}"
            })
        
        return tables
    
    def _analyze_document_structure(self, doc) -> Dict[str, Any]:
        """Analyze document structure and formatting"""
        
        structure = {
            'total_paragraphs': 0,
            'total_tables': 0,
            'sections': 0,
            'headings': [],
            'styles_used': set(),
            'has_headers': False,
            'has_footers': False
        }
        
        try:
            # Count paragraphs and analyze styles
            for paragraph in doc.paragraphs:
                structure['total_paragraphs'] += 1
                
                if paragraph.style:
                    style_name = paragraph.style.name
                    structure['styles_used'].add(style_name)
                    
                    # Detect headings
                    if 'heading' in style_name.lower() or 'title' in style_name.lower():
                        structure['headings'].append({
                            'text': paragraph.text[:100],  # First 100 chars
                            'style': style_name,
                            'level': self._extract_heading_level(style_name)
                        })
            
            # Count tables
            structure['total_tables'] = len(doc.tables)
            
            # Count sections
            structure['sections'] = len(doc.sections)
            
            # Check for headers/footers
            for section in doc.sections:
                if section.header and any(p.text.strip() for p in section.header.paragraphs):
                    structure['has_headers'] = True
                if section.footer and any(p.text.strip() for p in section.footer.paragraphs):
                    structure['has_footers'] = True
            
            # Convert set to list for JSON serialization
            structure['styles_used'] = list(structure['styles_used'])
        
        except Exception as e:
            structure['error'] = f"Structure analysis failed: {str(e)}"
        
        return structure
    
    def _extract_metadata(self, doc, filename: str, file_size: int) -> Dict[str, Any]:
        """Extract document metadata"""
        
        metadata = {
            'filename': filename,
            'size_bytes': file_size,
            'handler': self.handler_name
        }
        
        try:
            # Core properties
            core_props = doc.core_properties
            
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'category': core_props.category or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0
            })
        
        except Exception as e:
            metadata['metadata_error'] = f"Metadata extraction failed: {str(e)}"
        
        return metadata
    
    def _detect_images(self, doc) -> List[Dict[str, Any]]:
        """Detect images and shapes in document"""
        
        images = []
        
        try:
            # Basic implementation - count inline shapes
            image_count = 0
            
            # Check document parts for images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            if image_count > 0:
                images.append({
                    'total_images': image_count,
                    'note': 'Image extraction requires additional processing',
                    'detection_method': 'relationship_analysis'
                })
        
        except Exception as e:
            images.append({
                'error': f"Image detection failed: {str(e)}"
            })
        
        return images
    
    def _detect_table_header(self, table_data: List[List[str]]) -> bool:
        """Detect if table has header row"""
        
        if not table_data or len(table_data) < 2:
            return False
        
        first_row = table_data[0]
        second_row = table_data[1]
        
        # Simple heuristic: if first row has different characteristics than second
        first_row_avg_len = sum(len(cell) for cell in first_row) / len(first_row)
        second_row_avg_len = sum(len(cell) for cell in second_row) / len(second_row)
        
        # If first row is significantly shorter, likely a header
        return first_row_avg_len < second_row_avg_len * 0.7
    
    def _classify_table_type(self, table_data: List[List[str]]) -> str:
        """Classify table type based on content"""
        
        if not table_data:
            return 'unknown'
        
        # Join all table text for analysis
        table_text = ' '.join(' '.join(row) for row in table_data).lower()
        
        # Financial table indicators
        financial_indicators = ['revenue', 'income', 'profit', 'loss', 'assets', 'liabilities', '$', '%']
        if any(indicator in table_text for indicator in financial_indicators):
            return 'financial'
        
        # Data table indicators
        if any(char.isdigit() for char in table_text):
            return 'data'
        
        return 'text'
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        
        style_lower = style_name.lower()
        
        # Extract number from heading style
        for i in range(1, 10):
            if f'heading {i}' in style_lower or f'heading{i}' in style_lower:
                return i
        
        # Default level
        if 'title' in style_lower:
            return 1
        elif 'heading' in style_lower:
            return 2
        
        return 0
    
    def get_processing_estimate(self, uploaded_file) -> Dict[str, Any]:
        """Get processing time estimate for DOCX"""
        
        file_size_mb = uploaded_file.size / (1024 * 1024)
        estimated_time = self._estimate_processing_time(uploaded_file.size)
        
        return {
            'estimated_time_seconds': estimated_time,
            'file_size_mb': round(file_size_mb, 2),
            'complexity': 'high' if file_size_mb > 3 else 'medium' if file_size_mb > 1 else 'low',
            'docx_library_available': self.docx_available
        }